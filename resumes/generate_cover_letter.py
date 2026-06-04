import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml

def build_cover_letter():
    doc = Document()

    # Set Margins to 1 inch for standard professional page layout
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Text Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5) # Slightly smaller font size to ensure a perfect 1-page fit
    font.color.rgb = RGBColor(43, 42, 40) # Muted charcoal body text

    # Header Styling
    # Candidate Name
    p_name = doc.add_paragraph()
    run_name = p_name.add_run("LAILA MOHAMED FIKRY")
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(16)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(122, 0, 93) # Deep magenta theme color
    p_name.paragraph_format.space_after = Pt(2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Candidate Contact Details
    p_contact = doc.add_paragraph()
    contact_text = "Badr City, Cairo, Egypt  |  +20 121 021 2792  |  laila.mohamed.fikry@gmail.com\nLinkedIn: linkedin.com/in/laila-mohamed23/  |  GitHub: github.com/laila2005"
    run_contact = p_contact.add_run(contact_text)
    run_contact.font.name = 'Calibri'
    run_contact.font.size = Pt(9)
    run_contact.font.color.rgb = RGBColor(113, 111, 108) # Muted grey
    p_contact.paragraph_format.space_after = Pt(8)

    # Accent divider line
    p_hr = doc.add_paragraph()
    p_hr_border = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="7A005D"/>'
                            r'</w:pBdr>')
    p_hr._p.get_or_add_pPr().append(p_hr_border)
    p_hr.paragraph_format.space_after = Pt(12)

    # Date
    p_date = doc.add_paragraph()
    run_date = p_date.add_run("June 1, 2026")
    run_date.font.bold = True
    p_date.paragraph_format.space_after = Pt(8)

    # Recipient Info
    p_recip = doc.add_paragraph()
    recip_text = "Hiring Committee\nAnalyticsmart Consulting Inc.\nToronto, ON, Canada"
    run_recip = p_recip.add_run(recip_text)
    p_recip.paragraph_format.space_after = Pt(12)

    # Subject
    p_subject = doc.add_paragraph()
    run_subject = p_subject.add_run("Subject: Application for the Junior Back-End Developer Position")
    run_subject.font.bold = True
    run_subject.font.color.rgb = RGBColor(122, 0, 93)
    p_subject.paragraph_format.space_after = Pt(12)

    # Salutation
    p_salut = doc.add_paragraph()
    p_salut.add_run("Dear Hiring Committee,")
    p_salut.paragraph_format.space_after = Pt(8)

    # 4 Paragraphs body text to perfectly fit one page
    paragraphs = [
        "I am writing to express my strong enthusiasm for the Junior Back-End Developer position at Analyticsmart Consulting Inc., as advertised on LinkedIn. As a Year-3 Computer Science student at El Sewedy University of Technology with production backend engineering and deep learning deployment experience, I am excited to contribute immediately to your high-growth analytics, CRM, and AI-driven platforms.",
        
        "Despite being early in my career, I bring enterprise-level backend experience that matches your technical requirements. As Lead Software Engineer at LM Tech Solutions, I architected the RMS 3.0 Enterprise IoT Platform. I built a fault-tolerant backend service in C# and ASP.NET to manage concurrent, real-time data ingestion across Modbus, HTTP, and SNMP protocols for national enterprises, including the Egyptian Natural Gas Holding Company (GASCO) and the Ministry of Interior, while optimizing SQL Server database specifications.",
        
        "Furthermore, I possess a strong foundation in Node.js and RESTful API integration, alongside extensive hands-on experience deploying AI solutions. I architected Inqaz-app, an end-to-end emergency AI response system that utilizes deep learning computer vision to ingest live camera footage, perform real-time incident severity classification, and coordinate automated dispatch alerts. I am highly comfortable using advanced AI-assisted development tools like Claude Code and Gemini Antigravity, which allows me to deliver clean, highly optimized backend solutions at an accelerated pace.",
        
        "Although I am completing my CS degree (graduation expected 2027), I have structured my academic schedule to ensure full-time operational availability. Having successfully managed intense full-stack training programs (such as ALX Africa) and remote contract development roles, I am fully equipped to collaborate seamlessly with your international teams. Thank you for your time and consideration; I look forward to the opportunity to discuss how my unique skill set can add value to Analyticsmart."
    ]

    for p_text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)
        p.add_run(p_text)

    # Sign-off
    p_sign = doc.add_paragraph()
    p_sign.add_run("Sincerely,\n\nLaila Mohamed Fikry")
    p_sign.paragraph_format.space_after = Pt(8)

    output_path = r"G:\lolo\job_search\resumes\Laila_Mohamed_Fikry_Cover_Letter_Analyticsmart.docx"
    doc.save(output_path)
    print(f"Success: One-page Cover letter saved to {output_path}")

if __name__ == "__main__":
    build_cover_letter()
